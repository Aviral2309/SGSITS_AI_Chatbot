from urllib.parse import urlparse
import requests
from langchain_core.documents import Document

## fetching url
file_path = 'pages_name.txt'

lists=[]
with open(file_path, 'r') as file:
  for line in file:
    url = line.strip()
    lists.append(url)
  
  
from bs4 import BeautifulSoup
import re

def parse_html(html_content):
  soup = BeautifulSoup(html_content, 'html.parser')
  return soup.get_text()

def fetch_webpage(url):
  response = requests.get(url)
  return response.text

def clean_text(text):
  # remove extra whitespace and new lines
  text = ' '.join(text.split())
  # remove any remaining unwanted characters or patterns (eg: urls, special char)
  text = re.sub(r'http\S+','',text) #remove urls
  text = re.sub(r'[^a-zA-Z0-9\s.,!?\'"]+', '', text) # remove special character except punctuation
  return text


from docx import Document

def write_to_docx(text_content, output_filename):
  doc = Document() #create a new document object
  
  # add a heading
  doc.add_heading('Document Title', level =1)
  
  doc.add_paragraph(text_content)
  
  # save the document
  doc.save(output_filename)
  
for url in lists:
  html_content = fetch_webpage(url)
  page_text = parse_html(html_content)
  final_text = clean_text(page_text)
  output_filename = url.split('/')[-1] +'.docx'
  
  write_to_docx(final_text, output_filename)